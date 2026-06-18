"""Generate thinking_v2-0.docx — design document for OmniSME v2.0."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading("OmniSME AI Video Generator — Design Document v2.0", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Version: 2.0  |  Date: 2026-06-08  |  Status: Deployed")
doc.add_paragraph("Supersedes: thinking_v1-0.docx (v1.0 archived in releases/v1.0/)")
doc.add_paragraph("")

# ── 1. Change Summary ──────────────────────────────────────────────────────────
doc.add_heading("1. Changes from v1.0", level=1)
changes = [
    ("YouTube Source Footage in Video",
     "v1.0 produced an avatar-only video on a plain background. "
     "v2.0 downloads the first 2 minutes of each provided YouTube URL (via yt-dlp) "
     "and uses those clips as full-screen background footage. The HeyGen avatar "
     "narrates from the bottom-right corner in a picture-in-picture (PiP) overlay. "
     "An 'Avatar Only' fallback layout is still available via the Advanced Options."),
    ("Like / Regenerate controls",
     "After a video is generated, three action buttons appear below the player: "
     "'Like this video' (saves to the session's liked-videos list in the sidebar), "
     "'Regenerate' (re-runs the full pipeline with the same topic and URLs), and "
     "'New Topic' (clears all state for a fresh generation)."),
    ("New Topic shortcut",
     "A '+ New Topic' button is now present both at the top of the form and below "
     "the result panel, so the user can start over at any point without manually "
     "clearing inputs."),
    ("Version rollback support",
     "The previous working code is preserved in releases/v1.0/. "
     "To rollback: copy releases/v1.0/* back to the project root."),
]
for heading, body in changes:
    p = doc.add_paragraph(style="List Number")
    p.add_run(heading).bold = True
    doc.add_paragraph(body)

# ── 2. Architecture ────────────────────────────────────────────────────────────
doc.add_heading("2. Updated Architecture (v2.0)", level=1)
doc.add_paragraph("The pipeline now runs in 7 stages:")
stages = [
    ("1 — Input Collection",
     "Streamlit UI: topic text box, YouTube URL text area, avatar/voice dropdowns, layout selector."),
    ("2 — Transcript Extraction (LangChain)",
     "YoutubeLoader pulls transcripts from each URL. FAISS vector store built with "
     "OpenAI text-embedding-3-small embeddings."),
    ("3 — Script Generation (LangChain + GPT-4o)",
     "RetrievalQA retrieves relevant chunks; LLMChain (gpt-4o) formats into a "
     "structured presenter script (Hook → Intro → Key Points → Examples → Summary, ≤1,500 words)."),
    ("4 — YouTube Clip Download (yt-dlp)  [NEW]",
     "First 120 s of each YouTube URL downloaded to a temp folder using yt-dlp. "
     "Skipped silently if PiP layout is not selected."),
    ("5 — HeyGen Avatar Video",
     "Script split into ≤1,400-char slides, submitted to HeyGen v2 API. "
     "Polled until status = 'completed'. Avatar MP4 downloaded to ~/output/."),
    ("6 — Video Merge / PiP (moviepy)  [NEW]",
     "YouTube clips concatenated to match avatar duration (looped if needed). "
     "Avatar resized to 28 % width and composited in the bottom-right corner. "
     "YouTube audio muted; avatar narration audio preserved. Output: ~/output/<topic>_<ts>.mp4."),
    ("7 — Results & Controls  [NEW]",
     "Streamlit displays the merged video with Like / Regenerate / New Topic buttons. "
     "Liked videos are tracked in session state and listed in the sidebar."),
]
for label, body in stages:
    p = doc.add_paragraph(style="List Number")
    p.add_run(label).bold = True
    doc.add_paragraph(body)

# ── 3. Technology Stack ────────────────────────────────────────────────────────
doc.add_heading("3. Technology Stack (v2.0)", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Component", "Library / Service", "Version"
rows = [
    ("Frontend", "Streamlit", ">=1.35"),
    ("Transcript Extraction", "langchain-community YoutubeLoader", ">=0.2"),
    ("LLM Orchestration", "LangChain core + FAISS", ">=0.2"),
    ("LLM Backend", "OpenAI GPT-4o (gpt-4o)", "latest"),
    ("Embeddings", "OpenAI text-embedding-3-small", "latest"),
    ("YouTube Clip Download", "yt-dlp  [NEW v2.0]", ">=2024.1.0"),
    ("Video Composition", "moviepy (PiP merge)  [NEW v2.0]", ">=1.0.3,<2.0"),
    ("FFmpeg binaries", "imageio-ffmpeg  [NEW v2.0]", ">=0.4.9"),
    ("Avatar Video", "HeyGen REST API v2", "v2"),
    ("HTTP Client", "requests", ">=2.31"),
    ("Env Management", "python-dotenv", ">=1.0"),
]
for comp, lib, ver in rows:
    r = table.add_row().cells
    r[0].text, r[1].text, r[2].text = comp, lib, ver
doc.add_paragraph("")

# ── 4. File Structure ──────────────────────────────────────────────────────────
doc.add_heading("4. File & Folder Structure (v2.0)", level=1)
code = (
    "OmniSME_2/\n"
    "├── VERSION                   # '2.0'\n"
    "├── app.py                    # Streamlit UI (v2.0)\n"
    "├── requirements.txt          # Pinned dependencies (v2.0)\n"
    "├── src/\n"
    "│   ├── __init__.py\n"
    "│   ├── youtube_processor.py  # LangChain YoutubeLoader\n"
    "│   ├── script_generator.py   # FAISS + GPT-4o script builder\n"
    "│   ├── heygen_client.py      # HeyGen API wrapper\n"
    "│   ├── youtube_downloader.py # yt-dlp clip downloader  [NEW]\n"
    "│   └── video_merger.py       # moviepy PiP compositor  [NEW]\n"
    "└── releases/\n"
    "    └── v1.0/                 # Rollback snapshot\n"
    "        ├── VERSION\n"
    "        ├── app.py\n"
    "        ├── requirements.txt\n"
    "        └── src/\n"
    "\n"
    "~/ENV/.env        # HEYGEN_API_KEY + OPENAI_API_KEY (not modified)\n"
    "~/output/         # Generated MP4 files"
)
p = doc.add_paragraph()
p.add_run(code).font.name = "Courier New"

# ── 5. Component Details ───────────────────────────────────────────────────────
doc.add_heading("5. New Component Details", level=1)

doc.add_heading("5.1  src/youtube_downloader.py  [NEW]", level=2)
doc.add_paragraph(
    "Uses yt-dlp to download the first 120 seconds of each YouTube URL into a "
    "temporary directory. Format: best MP4 at ≤720p. Uses yt_dlp.utils.download_range_func "
    "to limit download length. Falls back gracefully if a video is unavailable or "
    "geo-restricted. Returns a list of local Path objects for the merger."
)

doc.add_heading("5.2  src/video_merger.py  [NEW]", level=2)
doc.add_paragraph(
    "Loads the HeyGen avatar MP4 and each YouTube clip with moviepy. "
    "Concatenates YouTube clips into a background track matching the avatar duration "
    "(loops the last clip if needed, all YouTube audio muted). "
    "Composites avatar at 28 % width in the bottom-right corner with 20 px padding. "
    "Preserves avatar audio as the final track. "
    "Writes the merged MP4 to ~/output/ via libx264 / aac."
)

doc.add_heading("5.3  app.py — Result Panel  [UPDATED]", level=2)
result_items = [
    "👍 Like this video: marks the result as liked; adds to sidebar liked-videos list.",
    "🔄 Regenerate: sets do_regenerate=True in session_state, clears video_result, reruns — "
    "pipeline executes again with the same topic and URLs (new GPT-4o temperature sample).",
    "＋ New Topic: resets all session_state to defaults and reruns, presenting a blank form.",
]
for item in result_items:
    doc.add_paragraph(item, style="List Bullet")

# ── 6. Rollback Procedure ──────────────────────────────────────────────────────
doc.add_heading("6. Rollback to v1.0", level=1)
doc.add_paragraph(
    "The previous release is archived in releases/v1.0/. To restore it:"
)
rollback = [
    "cp releases/v1.0/app.py .",
    "cp releases/v1.0/requirements.txt .",
    "cp releases/v1.0/src/* src/",
    "echo 'v1.0' > VERSION",
    "pip install -r requirements.txt",
    "streamlit run app.py",
]
for cmd in rollback:
    p = doc.add_paragraph()
    p.add_run(cmd).font.name = "Courier New"

# ── 7. Environment Variables ───────────────────────────────────────────────────
doc.add_heading("7. Environment Variables (unchanged)", level=1)
env_code = (
    "HEYGEN_API_KEY=<your_heygen_api_key>\n"
    "OPENAI_API_KEY=<your_openai_api_key>   # used by LangChain GPT-4o + embeddings"
)
p2 = doc.add_paragraph()
p2.add_run(env_code).font.name = "Courier New"

# ── 8. Limitations ────────────────────────────────────────────────────────────
doc.add_heading("8. Limitations & Assumptions (v2.0)", level=1)
lims = [
    "YouTube clips are limited to the first 2 minutes per URL to keep download times manageable.",
    "yt-dlp requires network access; geo-restricted or members-only videos will be skipped.",
    "moviepy requires ffmpeg (provided by imageio-ffmpeg). First run may download binaries.",
    "Liked-videos list is session-only (in-memory). It resets when the Streamlit session ends.",
    "PiP merging adds 1–3 minutes of local processing time after HeyGen rendering.",
    "YouTube videos must have transcripts/captions for the script generation stage.",
    "HeyGen free-tier caps avatar video duration at ~5 minutes.",
]
for lim in lims:
    doc.add_paragraph(lim, style="List Bullet")

# ── 9. Deployment ─────────────────────────────────────────────────────────────
doc.add_heading("9. Deployment Steps (v2.0)", level=1)
deploy = [
    "Confirm ~/ENV/.env contains HEYGEN_API_KEY and OPENAI_API_KEY.",
    "cd OmniSME_2/",
    "pip install -r requirements.txt   # installs yt-dlp, moviepy, imageio-ffmpeg",
    "streamlit run app.py",
    "Open http://localhost:8501 in a browser.",
    "Enter topic, paste YouTube URLs, select avatar/voice and layout.",
    "Click 'Generate Video'. Use Like / Regenerate / New Topic as needed.",
]
for i, d in enumerate(deploy, 1):
    doc.add_paragraph(f"{i}. {d}")

# ── Save ───────────────────────────────────────────────────────────────────────
out = "/Users/hardikraval/Documents/Trainings/Claude/TestProject/OmniSME_2/thinking_v2-0.docx"
doc.save(out)
print(f"Document saved: {out}")
