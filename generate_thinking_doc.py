"""Script to generate thinking_v1-0.docx design document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Title ---
title = doc.add_heading("OmniSME AI Video Generator — Design Document", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Version: 1.0  |  Date: 2026-06-08  |  Status: Pending Approval")
doc.add_paragraph("")

# ─── 1. Project Overview ───────────────────────────────────────────────────
doc.add_heading("1. Project Overview", level=1)
doc.add_paragraph(
    "Build a web application using Python Streamlit that allows users to:"
)
items = [
    "Enter a topic and a list of YouTube video URLs as source material.",
    "Automatically extract and analyse content from those YouTube videos "
    "using the LangChain framework.",
    "Generate a professional, topic-focused video narrated by a HeyGen "
    "AI avatar that synthesises only the provided YouTube sources.",
    "Save the rendered video to ~/output/ on the local laptop.",
]
for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

# ─── 2. Architecture ───────────────────────────────────────────────────────
doc.add_heading("2. High-Level Architecture", level=1)
doc.add_paragraph(
    "The pipeline runs end-to-end in four stages:"
)
stages = [
    ("Stage 1 — Input Collection",
     "Streamlit UI collects the topic string and one or more YouTube URLs "
     "from the user."),
    ("Stage 2 — YouTube Content Extraction (LangChain)",
     "YoutubeLoader (langchain_community) pulls timestamped transcripts. "
     "A LangChain RetrievalQA chain with an in-memory FAISS vector store "
     "retrieves only the segments most relevant to the requested topic."),
    ("Stage 3 — Script Generation (LangChain + LLM)",
     "A LangChain LLMChain (using OpenAI GPT-4o) transforms the retrieved "
     "segments into a structured, presenter-style video script: introduction, "
     "3–5 key sections, and a closing summary."),
    ("Stage 4 — HeyGen Video Generation",
     "The script is sent to the HeyGen v2 REST API. The app selects an avatar "
     "and voice, submits the generation request, polls until completion, "
     "downloads the MP4, and writes it to ~/output/."),
]
for heading, body in stages:
    doc.add_paragraph(heading, style="List Number")
    doc.add_paragraph(body)

# ─── 3. Technology Stack ───────────────────────────────────────────────────
doc.add_heading("3. Technology Stack", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Component"
hdr[1].text = "Library / Service"
hdr[2].text = "Version Pinned"
rows = [
    ("Frontend", "Streamlit", ">=1.35"),
    ("Content Extraction", "langchain-community YoutubeLoader", ">=0.2"),
    ("LLM Orchestration", "LangChain core + FAISS", ">=0.2"),
    ("LLM Backend", "OpenAI GPT-4o (gpt-4o)", "latest"),
    ("Video Generation", "HeyGen REST API v2", "v2"),
    ("HTTP Client", "requests", ">=2.31"),
    ("Env Management", "python-dotenv", ">=1.0"),
    ("Document Generation", "python-docx", ">=1.1"),
    ("Transcript Download", "youtube-transcript-api", ">=0.6"),
]
for comp, lib, ver in rows:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = lib
    row[2].text = ver

doc.add_paragraph("")

# ─── 4. File Structure ─────────────────────────────────────────────────────
doc.add_heading("4. File & Folder Structure", level=1)
doc.add_paragraph(
    "All source files live inside the project directory. Config and output "
    "use home-directory paths to avoid checking in secrets."
)
code = (
    "OmniSME_2/\n"
    "├── app.py                  # Streamlit entry point\n"
    "├── requirements.txt        # Pinned dependencies\n"
    "└── src/\n"
    "    ├── __init__.py\n"
    "    ├── youtube_processor.py  # YouTube transcript extraction\n"
    "    ├── script_generator.py   # LangChain script builder\n"
    "    └── heygen_client.py      # HeyGen API wrapper\n"
    "\n"
    "~/ENV/.env        # HeyGen + LLM API keys (pre-existing, not modified)\n"
    "~/output/         # Generated MP4 files (pre-existing)"
)
p = doc.add_paragraph()
p.add_run(code).font.name = "Courier New"

# ─── 5. Component Details ──────────────────────────────────────────────────
doc.add_heading("5. Component Details", level=1)

doc.add_heading("5.1  app.py — Streamlit UI", level=2)
ui_items = [
    "Title bar and project branding.",
    "Text input: Topic (e.g. 'Cisco QoS best practices').",
    "Text area: YouTube URLs — one per line.",
    "Expander: Advanced options (avatar ID, voice ID, video duration hint).",
    "Generate button → triggers pipeline with st.spinner progress updates.",
    "On success: displays embedded video player and a Download button.",
    "Error messages surfaced via st.error with actionable hints.",
]
for i in ui_items:
    doc.add_paragraph(i, style="List Bullet")

doc.add_heading("5.2  src/youtube_processor.py", level=2)
doc.add_paragraph(
    "Uses langchain_community.document_loaders.YoutubeLoader to pull "
    "transcripts for each URL. Falls back to auto-generated captions when "
    "manual ones are unavailable. Returns a list of LangChain Document "
    "objects. Raises a descriptive error if a video has no transcript."
)

doc.add_heading("5.3  src/script_generator.py", level=2)
doc.add_paragraph(
    "Embeds the documents into a FAISS in-memory vector store "
    "(langchain.vectorstores.FAISS) using OpenAI embeddings (text-embedding-3-small). "
    "A RetrievalQA chain retrieves the top-k relevant chunks for the topic. "
    "A subsequent LLMChain (ChatOpenAI, model=gpt-4o) formats the retrieved content "
    "into a structured presenter script with sections: Hook, Background, Key Points (3–5), "
    "Demonstration / Examples, and Summary. "
    "The total script is kept under 1,500 words to stay within HeyGen "
    "avatar video length limits."
)

doc.add_heading("5.4  src/heygen_client.py", level=2)
heygen_details = [
    "list_avatars() — GET /v2/avatars, returns id/name pairs.",
    "list_voices() — GET /v2/voices, filtered to English.",
    "create_video(script, avatar_id, voice_id) — POST /v2/video/generate.",
    "poll_status(video_id, timeout=600) — polls GET /v1/video_status.get "
    "every 10 s until status is 'completed' or timeout exceeded.",
    "download_video(video_url, topic) — streams MP4 to "
    "~/output/<sanitised_topic>_<timestamp>.mp4.",
]
for d in heygen_details:
    doc.add_paragraph(d, style="List Bullet")

# ─── 6. HeyGen API Reference ───────────────────────────────────────────────
doc.add_heading("6. HeyGen API Integration", level=1)
table2 = doc.add_table(rows=1, cols=3)
table2.style = "Light Grid Accent 1"
h2 = table2.rows[0].cells
h2[0].text = "Endpoint"
h2[1].text = "Method"
h2[2].text = "Purpose"
api_rows = [
    ("/v2/avatars", "GET", "List available avatar IDs and names"),
    ("/v2/voices", "GET", "List available voice IDs"),
    ("/v2/video/generate", "POST", "Submit video generation request"),
    ("/v1/video_status.get", "GET", "Poll video generation status"),
    ("(video_url in response)", "GET", "Stream-download the rendered MP4"),
]
for ep, meth, purpose in api_rows:
    r = table2.add_row().cells
    r[0].text = ep
    r[1].text = meth
    r[2].text = purpose

doc.add_paragraph("")
doc.add_paragraph(
    "Authentication: X-Api-Key header using HEYGEN_API_KEY from "
    "~/ENV/.env. All requests use HTTPS."
)

# ─── 7. Environment Variables ──────────────────────────────────────────────
doc.add_heading("7. Environment Variables (~/ENV/.env)", level=1)
doc.add_paragraph(
    "The file at ~/ENV/.env must contain the keys below. The file is "
    "pre-existing and will NOT be recreated or overwritten by this project."
)
env_code = (
    "HEYGEN_API_KEY=<your_heygen_api_key>\n"
    "OPENAI_API_KEY=<your_openai_api_key>   # used by LangChain LLM + embeddings"
)
p2 = doc.add_paragraph()
p2.add_run(env_code).font.name = "Courier New"

# ─── 8. Data Flow ──────────────────────────────────────────────────────────
doc.add_heading("8. Detailed Data Flow", level=1)
steps = [
    "User submits topic + URLs via Streamlit.",
    "YouTubeProcessor.load(urls) → list of Document objects with transcript text.",
    "ScriptGenerator.build(docs, topic) → structured presenter script string.",
    "HeyGenClient.create_video(script, avatar_id, voice_id) → video_id.",
    "HeyGenClient.poll_status(video_id) → video_url when complete.",
    "HeyGenClient.download_video(video_url, topic) → local MP4 path in ~/output/.",
    "Streamlit displays video player with st.video(local_path) and download link.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

# ─── 9. Limitations & Assumptions ─────────────────────────────────────────
doc.add_heading("9. Limitations & Assumptions", level=1)
lims = [
    "YouTube videos must have auto-generated or manual captions. Videos "
    "without transcripts will be skipped with a warning.",
    "HeyGen free-tier accounts cap video duration at ~5 minutes. Paid plans "
    "support longer content.",
    "Script length is capped at 1,500 words to stay within HeyGen limits.",
    "An active OpenAI API key is required for LangChain LLM calls (GPT-4o) "
    "and embeddings (text-embedding-3-small) — this is separate from the HeyGen key.",
    "The app does not download or process raw video frames — it uses "
    "text transcripts only as source material.",
    "Internet connectivity is required throughout the pipeline.",
    "The generated video will strictly reference content from the provided "
    "YouTube URLs; no additional web sources are consulted.",
]
for lim in lims:
    doc.add_paragraph(lim, style="List Bullet")

# ─── 10. Deployment Plan ───────────────────────────────────────────────────
doc.add_heading("10. Deployment Steps", level=1)
deploy = [
    "Verify ~/ENV/.env contains HEYGEN_API_KEY and OPENAI_API_KEY.",
    "cd into OmniSME_2/ project directory.",
    "pip install -r requirements.txt",
    "streamlit run app.py",
    "Open browser at http://localhost:8501.",
    "Enter topic, paste YouTube URLs (one per line), click Generate.",
    "Wait for progress updates; video saved to ~/output/ on completion.",
]
for i, d in enumerate(deploy, 1):
    doc.add_paragraph(f"{i}. {d}")

# ─── Save ──────────────────────────────────────────────────────────────────
out = "/Users/hardikraval/Documents/Trainings/Claude/TestProject/OmniSME_2/thinking_v1-0.docx"
doc.save(out)
print(f"Document saved: {out}")
